"""
parsers/resume_parser.py
Extracts raw text from PDF, DOCX, TXT, and Markdown files.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn

console = Console()


@dataclass
class RawResume:
    path: Path
    filename: str
    text: str
    file_type: str
    char_count: int
    error: Optional[str] = None

    @property
    def is_valid(self) -> bool:
        return self.error is None and self.char_count > 100


class ResumeParser:
    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def parse_all(self, files: list[Path]) -> list[RawResume]:
        results = []
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Parsing resumes[/bold blue]"),
            BarColumn(),
            TaskProgressColumn(),
            TextColumn("{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("", total=len(files))
            for f in files:
                progress.update(task, description=f"[dim]{f.name}[/dim]")
                result = self._parse_file(f)
                results.append(result)
                if self.verbose:
                    if result.is_valid:
                        console.print(f"  [green]✓[/green] {f.name} ({result.char_count:,} chars)")
                    else:
                        console.print(f"  [red]✗[/red] {f.name}: {result.error}")
                progress.advance(task)

        valid = sum(1 for r in results if r.is_valid)
        console.print(f"[green]✓[/green] Parsed [bold]{valid}/{len(files)}[/bold] resumes successfully\n")
        return results

    def _parse_file(self, path: Path) -> RawResume:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                text = self._parse_pdf(path)
            elif ext in (".docx", ".doc"):
                text = self._parse_docx(path)
            elif ext in (".txt", ".md"):
                text = path.read_text(encoding="utf-8", errors="replace")
            else:
                return RawResume(path=path, filename=path.name, text="",
                                 file_type=ext, char_count=0,
                                 error=f"Unsupported format: {ext}")

            text = self._clean_text(text)
            return RawResume(
                path=path,
                filename=path.name,
                text=text,
                file_type=ext,
                char_count=len(text),
            )
        except Exception as e:
            return RawResume(path=path, filename=path.name, text="",
                             file_type=ext, char_count=0, error=str(e))

    def _parse_pdf(self, path: Path) -> str:
        # Try pdfplumber first (better layout), fallback to pypdf2
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            return "\n\n".join(pages)
        except ImportError:
            pass

        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            return "\n\n".join(pages)
        except ImportError:
            pass

        try:
            import pymupdf  # fitz
            doc = pymupdf.open(str(path))
            pages = [page.get_text() for page in doc]
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError(
                "No PDF library found. Install one:\n"
                "  pip install pdfplumber   (recommended)\n"
                "  pip install pypdf\n"
                "  pip install pymupdf"
            )

    def _parse_docx(self, path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

    def _clean_text(self, text: str) -> str:
        # Normalize whitespace
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        # Collapse excessive blank lines
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        # Remove null bytes and control chars (keep newlines and tabs)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text.strip()
